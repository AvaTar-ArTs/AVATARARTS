"""
File parsing service.
"""
import os
import logging
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


class FileParserService:
    """File parsing service."""

    SUPPORTED_EXTENSIONS = {
        'pdf': 'parse_pdf',
        'txt': 'parse_text',
        'md': 'parse_text',
        'docx': 'parse_docx',
        'doc': 'parse_docx',
        'xlsx': 'parse_excel',
        'xls': 'parse_excel',
        'csv': 'parse_csv',
    }

    def parse_file(self, file_path: str, filename: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Parse a file.

        Args:
            file_path: file path
            filename: file name

        Returns:
            (extracted text, error message)
        """
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        if ext not in self.SUPPORTED_EXTENSIONS:
            return None, f"Unsupported file format: {ext}"

        parser_method = getattr(self, self.SUPPORTED_EXTENSIONS[ext])

        try:
            content = parser_method(file_path)
            if not content:
                return None, "Unable to extract file content"
            return content, None
        except Exception as e:
            logger.error(f"File parsing failed: {e}")
            return None, str(e)

    def parse_pdf(self, file_path: str) -> Optional[str]:
        """Parse PDF file."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return '\n\n'.join(text_parts)

        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise

    def parse_text(self, file_path: str) -> Optional[str]:
        """Parse plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'big5', 'gb2312', 'latin-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            raise ValueError("Unable to detect file encoding")

        except Exception as e:
            logger.error(f"Text file parsing failed: {e}")
            raise

    def parse_docx(self, file_path: str) -> Optional[str]:
        """Parse Word document."""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Word parsing failed: {e}")
            raise

    def parse_excel(self, file_path: str) -> Optional[str]:
        """Parse Excel file."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, data_only=True)
            text_parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"## Sheet: {sheet_name}\n")

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_values):
                        text_parts.append(' | '.join(row_values))

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Excel parsing failed: {e}")
            raise

    def parse_csv(self, file_path: str) -> Optional[str]:
        """Parse CSV file."""
        try:
            import csv

            text_parts = []

            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'big5', 'gb2312', 'latin-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, newline='') as f:
                        reader = csv.reader(f)
                        for row in reader:
                            if any(row):
                                text_parts.append(' | '.join(row))
                    break
                except UnicodeDecodeError:
                    continue

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"CSV parsing failed: {e}")
            raise


# Global instance
_file_parser: Optional[FileParserService] = None


def get_file_parser() -> FileParserService:
    """Get file parsing service instance."""
    global _file_parser
    if _file_parser is None:
        _file_parser = FileParserService()
    return _file_parser
